from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize document
doc = Document()

# Heading
heading = doc.add_heading('BUG REPORT', 0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Summary / Title
title_p = doc.add_paragraph()
title_p.add_run('Title: ').bold = True
title_p.add_run('Authentication failure on login page with valid credentials - "Your email, password, IP address or location did not match"')

# Environment
env_p = doc.add_paragraph()
env_p.add_run('Environment: ').bold = True
env_p.add_run('Production (app.vwo.com)')

# Severity / Priority
sev_p = doc.add_paragraph()
sev_p.add_run('Severity: ').bold = True
sev_p.add_run('Critical (Authentication Bypass/Blocker)')

pri_p = doc.add_paragraph()
pri_p.add_run('Priority: ').bold = True
pri_p.add_run('High')

# Steps to Reproduce
doc.add_paragraph().add_run('Steps to Reproduce:').bold = True
doc.add_paragraph('1. Navigate to https://app.vwo.com/', style='List Number')
doc.add_paragraph('2. Enter valid Email ID: sample1234 in the input field.', style='List Number')
doc.add_paragraph('3. Enter the correct Password: Sample@1234 associated with the account.', style='List Number')
doc.add_paragraph('4. Click on the "Sign in" button.', style='List Number')

# Expected Result
exp_p = doc.add_paragraph()
exp_p.add_run('Expected Result: ').bold = True
exp_p.add_run('User should be successfully authenticated and redirected to the VWO Dashboard.')

# Actual Result
act_p = doc.add_paragraph()
act_p.add_run('Actual Result: ').bold = True
act_p.add_run('Authentication fails and a generic error message is displayed: "Your email, password, IP address or location did not match."')

# Evidence
doc.add_paragraph().add_run('Evidence:').bold = True
img_path = r"w:\The Testing Acedamy\Antigravity\My Projects\Project_01_VWO\Input\app.vwo.com Login_Error.png"
try:
    doc.add_picture(img_path, width=Inches(5.0))
except Exception as e:
    doc.add_paragraph(f'[Error adding image: {str(e)}]')

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run('--- End of Bug Report ---').italic = True

# Save Document
output_path = r"w:\The Testing Acedamy\Antigravity\My Projects\Project_01_VWO\Test_ Delivarables_Outcome\VWO_Bug_Login_Page.doc"
import os
if os.path.exists(output_path):
    try:
        os.remove(output_path)
    except:
        pass
        
doc.save(output_path)
print(f"Bug report successfully updated with valid credentials: {output_path}")
