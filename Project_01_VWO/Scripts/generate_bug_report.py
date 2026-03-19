from docx import Document

# Initialize document
doc = Document()
doc.add_heading('Bug Report: VWO Login Page', 0)

# Apply Bug Template
records = {
    'Title': 'Error message "You are not allowed to log in." displayed upon login with Valid Username and Valid Password',
    'Environment': 'app.vwo.com',
    'Severity': '[NEEDS CLARIFICATION] (Cannot determine from notes - likely High/Critical as login is blocked)',
    'Steps to Reproduce': '1. Navigate to app.vwo.com\n2. Enter a Valid Username\n3. Enter a Valid Password\n4. Click on the submit button',
    'Expected Result': 'Successful login and navigation to the dashboard [NEEDS CLARIFICATION on expected destination]',
    'Actual Result': 'An error message is coming stating exactly: "You are not allowed to log in."',
    'Evidence Required': '[NEEDS CLARIFICATION] Requesting screenshot of error, console logs, and specific Valid Username/Password used for the test.'
}

for key, value in records.items():
    p = doc.add_paragraph()
    p.add_run(f'{key}: ').bold = True
    if '\n' in value:
        doc.add_paragraph(value)
    else:
        p.add_run(value)

doc.add_heading('Notes Analyzed', level=2)
doc.add_paragraph(
    'What I have done is I have navigated to app.vwo.com. I have basically added an Valid Username and Valid Password. '
    'When I click on the submit button, I can also see that there is a remember me button, and there are social logins '
    'also available when I saw the overall application, which is app.vw.com. Then after that, what I have noticed is that '
    'whenever I click on submit button, there is an error message which is coming. That error message is "You are not allowed to log in."'
)

doc.add_heading('Anti-Hallucination Checks', level=2)
doc.add_paragraph('✅ Used only provided evidence.')
doc.add_paragraph('✅ Marked unknowns explicitly as [NEEDS CLARIFICATION].')
doc.add_paragraph('✅ Did not assume root cause or invent configurations.')

# Save Document
output_path = r"w:\The Testing Acedamy\Antigravity\My Projects\Project_01_VWO\Test_ Delivarables_Outcome\VWO_Bug_Login_Page.doc"
doc.save(output_path)
print(f"Doc Saved: {output_path}")
