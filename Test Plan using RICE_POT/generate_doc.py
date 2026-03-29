from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_plan():
    doc = Document()

    # Title
    title = doc.add_heading('Enterprise Master Test Plan: VWO Login Dashboard', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.add_run('Document Version: 1.0 | Status: Draft for Stakeholder Review').bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sections = [
        ("1. Test Plan ID", "TP-VWO-LOG-2026-001"),
        ("2. Instruction", "This Test Plan serves as the master blueprint for verifying the VWO Login Dashboard. All QA activities must adhere to the VWO Governance standards. Automated scripts must be integrated into the CI/CD pipeline, and manual edge-case testing must be documented in the JIRA Xray module. Strict adherence to the 2-second load time SLA and SOC2/GDPR compliance is mandatory."),
        ("3. Test Objectives", "* Security Validation: Ensure zero vulnerability in authentication flows including MFA, SSO, and Passkey.\n* Performance Benchmark: Verify the login page loads within 2 seconds under a load of 5,000 concurrent users.\n* Accessibility (ADA): Achieve 100% compliance with WCAG 2.1 Level AA standards for screen readers and keyboard-only users.\n* User Experience: Validate seamless transitions to the VWO core platform post-authentication with a target success rate of 95%."),
        ("4. Scope", "In-Scope:\n- Primary Email/Password Authentication (with 'Remember Me').\n- Multi-Factor Authentication (MFA) dependency and flow.\n- Enterprise SSO (SAML/OAuth) and Social Login (Google).\n- Password Management (Forgot Password, Reset Token, Password Strength).\n- Responsive UI verification across iOS, Android, and Desktop browsers.\n\nOut-of-Scope:\n- Core Dashboard features (post-login navigation only).\n- Third-party marketing tool backend configurations."),
        ("5. Test Strategy", "Testing Level: System Integration Testing (SIT) followed by User Acceptance Testing (UAT).\n\nTesting Types:\n- Functional: Boundary Value Analysis (BVA) for password complexity and email formats.\n- Automation: Selenium/Playwright scripts for regression; nightly executions.\n- Security: OWASP Top 10 scanning, brute-force rate limiting, and session hijacking prevention tests.\n- Performance: JMeter/K6 for load testing (2s SLA verification).\n- Accessibility: Axe-core and manual NVDA/JAWS screen reader testing."),
        ("6. Test Environment", "Browsers: Chrome (Latest), Firefox (Latest), Safari (Mac/iOS), Edge\nMobile Devices: iPhone 15 (Safari), Samsung S24 (Chrome), iPad Pro\nOperating Systems: Windows 11, macOS Sonoma, iOS 17, Android 14\nNetwork: 4G, 5G, Broadband, and Throttled 3G (for performance latency)"),
        ("7. Test Scenarios/Cases", "SC_LOG_01: Successful login with valid credentials and 'Remember Me' session persistence.\nSC_LOG_02: Verification of MFA token expiry and 'locked out' state after 5 failed attempts.\nSC_LOG_03: Redirect flow from Google Auth/Enterprise SAML back to VWO Dashboard.\nSC_LOG_04: Keyboard-only navigation (Tab + Enter) to trigger 'Sign In with Passkey'.\nSC_LOG_05: Page load time measurement during a 10% traffic spike simulation.\nSC_LOG_06: Password input visibility toggle (Eye Icon) validation."),
        ("8. Entry Criteria", "* PRD and UI Mockups are finalized and signed off.\n* Test Environment is provisioned with valid SSL/TLS certificates.\n* Smoke test of the build passes (Login page is accessible).\n* Test Data (Test Users, SSO IDs, MFA Tokens) is prepared."),
        ("9. Exit Criteria", "* 100% of Critical and High-priority Test Cases are passed.\n* Zero Open Defects with P0 (Critical) or P1 (High) severity.\n* Accessibility Score > 95% on Lighthouse/Axe tools.\n* Performance reports confirm load time < 2 seconds."),
        ("10. Risk/Mitigation", "Risk: MFA Delay | Mitigation: Work with DevOps to ensure SMS/Email gateway latency is < 5s.\nRisk: SSO Failure | Mitigation: Implement a fallback to Email/Password login for emergency access.\nRisk: Browser Compatibility | Mitigation: Use SauceLabs/BrowserStack for Cross-Browser automated regression.\nRisk: Rate Limiting | Mitigation: Whitelist QA IPs during performance testing to avoid false negatives."),
        ("11. Deliverables", "* Complete Master Test Plan (This Document).\n* Automation Test Suite (Script Repository).\n* Defect Logs (JIRA Dashboard).\n* Test Summary Report (TSR) with Sign-off.\n* Accessibility Compliance Certificate."),
        ("12. Schedule", "* Requirement Analysis: 2 Days\n* Test Case Design: 3 Days\n* Execution (SIT): 5 Days\n* UAT & Performance: 3 Days\n* Final Sign-off: 1 Day"),
        ("13. Role & Responsibility", "QA Manager: Approver (Accountable for overall quality).\nQA Automation Lead: Responsible for script development and execution.\nProduct Manager: Informed (Signs off on functional requirements).\nDevOps/Security Team: Consultation on SSO and MFA infra."),
        ("14. Summary of Analysis", "Specifically addresses the Security Matrix (MFA/SSO), ADA Compliance (WCAG 2.1 AA), and Performance targets (2s load time) identified in the VWO Login Dashboard PRD.")
    ]

    for title, content in sections:
        h = doc.add_heading(title, level=1)
        p = doc.add_paragraph(content)

    doc.save('VWO_Login_Dashboard_Test_Plan.docx')

if __name__ == "__main__":
    create_test_plan()
